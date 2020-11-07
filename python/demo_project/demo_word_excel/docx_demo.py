from docx import Document
from docx.shared import Inches, Pt

import os

# 官方文档 https://python-docx.readthedocs.io/en/latest/

# 写入 docx 文档
def write_docx():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')

script_dir = os.path.dirname(os.path.realpath(__file__))

def read_docx():
    # read_demo = os.path.join(script_dir, '文档样例.docx')
    read_demo = os.path.join(script_dir, 'new_demo.docx')
    document = Document(read_demo)
    # document.save(os.path.join(script_dir, 'new_demo.docx'))

    paras = document.paragraphs
    for para in paras:
        print(para)
        print("style:", para.style)
        font = para.style.font
        print(font.name, font.size, font.bold, font.italic)
        for runItem in para.runs:
            # print("run:", runItem.text)
            if runItem.text.startswith('所有的费用'):
                print('change size')
                para.style = 'Heading 2'
                para.style.font.size=Pt(24)
        print(para.text)

    # document.save(os.path.join(script_dir, 'new_demo.docx'))

if __name__ == '__main__':
    # write_docx()
    read_docx()